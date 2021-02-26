# Data Extraction #####
import re
import docx2txt
from docx import Document
from nltk.stem import PorterStemmer
from nltk.stem import LancasterStemmer
from docx import Document
###################### complete Qui
from docx import Document
from tkinter import *
from tkinter import filedialog
import tkinter.font as font
from tkinter import messagebox
import os



############################ for generaliazation #######################################

text_realted_dictionary = {}
Main_heading_list = ["hobbies","Skill","education","work","experience","certification","affiliation","projects",
                    "researches","publication","activities","information","interests","career","qualification","academic","expertise","objectives",
                     "training","volunteering","languages","studies","miscellaneous","education","employment","profile",
                     "summary","jobs","expertise","Competencies","Operating","ICT"]
languages_used = ["Spanish","English","Hindi","Arabic","Portuguese","Bengali","Russian","Japanese","Punjabi","German",
                  "Javanese","Malay","Telugu","Vietnamese","Korean","French","Marathi","Tamil","Urdu","Turkish","Italian",
                  "Thai","Gujarati","Jin","Southern Min" ,"Persian","Polish","Pashto","Kannada","Xiang","Malayalam",
                  "Sundanese","Hausa","Odia","Burmese","Hakka","Ukrainian","Bhojpuri","Tagalog","Yoruba","Maithili",
                  "Uzbek","Sindhi","Amharic","Fula","Romanian","Oromo","Igbo","Azerbaijani","Awadhi","Gan","Cebuano"]
Main_heading_after_stemming = []
Final_heading = []
Data_related_to_headings =[]

templete_dictionary={}
extraction_dictionary={}

####################### Data Ectraction varaibles ############################################

Mainheadertittle = ""                                   # Main header heading is present in this varaiable
MainFootertittle = ""                                   # Main footer heading is present in this varaiable 

####################### Docx2 Varaiables #############################################

fullText_doc2 = []                                      # Main text List of Docx2

####################### Python-docx function ##################################################\
heading_length_extraction =3
heading_length_templete = 3
templete =""
Extraction =""
data_extraction_heading = []
templete_heading = []
 #####################################################################################################

    
def Finding_header_():
    document = Document("Normal/8.docx")                #Reading the file viva python-docx
    section = document.sections[0]
    header =section.header                              # reading the header from the docx file
    for paragraph in header.paragraphs:
        return paragraph.text
        
def Finding_footer_():
    document = Document("Normal/8.docx")                #Reading the file viva python-docx
    section = document.sections[0]
    footer =section.footer                              # reading the footer from the docx file
    for paragraph in footer.paragraphs:
        return paragraph.text

############################### General Coding ################################################

def Words_adjustments(word):                                # words_adjustments is basically removing the tab and spaces 
    word_list = word.split()
    new_word = ""
    for i in word_list:
        new_word = new_word + i + " "
    return (new_word.strip())

def matching_function(word,i):                             # This function is matching the heads and words  
    Main_heading_after_stemming = Main_heading_list_stemming_function(Main_heading_list)
    if word.lower() in Main_heading_after_stemming:
        Final_heading.append(i)
        Data_related_to_headings.append(0)
        return True
    
        
        
def heading_finding_pre_steps(lists,heading_length):                   # Limited the searching of headings Step-1 in which len 5 is included

    for i in lists:
        Second_list = i.split()                         # second_list is just Heading finding in limited searches
  
        if len(Second_list) <= heading_length:
            if len(Second_list) != 0:
                for j in Second_list:
                    word = normal_stemming_function(j)
                    Flag = matching_function(word,i)
                    if Flag:
                        break
                if Flag:                                              #Checking that if it is not heading then it means
                    continue                                          #that is going to be detial of certain heading
                else:
                    Data_related_to_headings.append(i)

        else:
            Data_related_to_headings.append(i)                      #if it is not going to len related to heading then
                                                                    #it is detail of cetain heading
    return Final_heading
            
            
def Main_heading_list_stemming_function(lists):           # Main Heading by defult stemmming process 
    porter = PorterStemmer()
    lancaster=LancasterStemmer()
    length = len(lists)
    counter = 0
    while(length > counter):
        Main_heading_after_stemming.append(porter.stem(lists[counter]))
        counter = counter + 1
    return (Main_heading_after_stemming)


def normal_stemming_function(word):                     # Words stemming 
    porter = PorterStemmer()
    lancaster=LancasterStemmer()
    return porter.stem(word)

def now_printing_all_the_data_related_it(Data_related_to_headings,Final_heading):
    for i in Final_heading:
        text_realted_dictionary[i] = ""
        
    Text=[] 
    counter = 0
    heading=""
    text_docx=""
    for i in Data_related_to_headings:

        if i ==0:  
            if counter > 0:
                Text.append(text_docx)
                text_docx =""
            counter = counter + 1;
        else:
            text_docx = text_docx + i + "\n"
    Text.append(text_docx)
    return Text

#################################################### Major Calling Function ############################################

def main_function(path,heading_length):

    text_realted_dictionary = {}
    Text = []
    heading=[]
    del  Final_heading[:]
    del Main_heading_after_stemming[:]
    del Data_related_to_headings[:]
    del fullText_doc2[:]
    MY_TEXT = docx2txt.process(path)              #Reading the file viva docx2txt libaray 
    Text =MY_TEXT.split("\n")                                #Split the whole text at new line character

    Mainheadertittle = Finding_header_()                    # checking the header is avaiable or not 
    MainFootertittle=  Finding_footer_()                    # checking the footer is avaiable or not
    
                 # Bring the data is proper fromat before doing process on it 
    for i in Text:
        if len(i) >1:
            if i != Mainheadertittle:
                if i != MainFootertittle:
                    fullText_doc2.append(Words_adjustments(i))   # Words_adjustments is Basically removing tab and space in the end
    heading = heading_finding_pre_steps(fullText_doc2,heading_length)
    without_duplicate = set(Final_heading)
    Text = now_printing_all_the_data_related_it(Data_related_to_headings,Final_heading)

    counter=0;
    for i in Final_heading:
        text_realted_dictionary[i.lower()] = Text[counter]
        counter = counter + 1
    return heading[:],text_realted_dictionary


 #####################################################################################################
    
def matching_heading_between_both(data_extraction_heading,templete_heading):
    data_extraction_heading=[x.lower() for x in data_extraction_heading]
    templete_heading=[x.lower() for x in templete_heading]

    matching_heading =[]
    dic_heading = []

    connection_word=["the","and","an","a","is","cont","of"]
    flag = False

    for heading_extraction in data_extraction_heading:
        h_e_split=heading_extraction.split()
        for i in h_e_split:
            if i in connection_word:
                continue 
            else:
                for heading_cv in templete_heading:
                    h_c_split = heading_cv.split()
                    for j in h_c_split:
                        if normal_stemming_function(j) == normal_stemming_function(i):
                            matching_heading.append(heading_cv.lower())
                            dic_heading.append(heading_extraction.lower())
                            flag = True
                            break

                    if flag == True:
                        break
                if flag == True:
                    flag = False
                    break
#     print("The matching hae matching_heading")
    return matching_heading,dic_heading

###### Main Insertion Wala Part Tayyab ##########
def add_para(data,doc):
    return doc.add_paragraph(data) # adding the paragraph at the end

def move_para(document,matching_heading,dic_heading):
    for paragraph in document.paragraphs: #Iterating through each paragraph and checking when the paragraph text is equal
                                            # to headings of template this means now there is heading in paragraph variable
        if (paragraph.text).lower() in matching_heading:
            indexing = matching_heading.index(paragraph.text.lower())
            para = add_para(extraction_dictionary[dic_heading[indexing].lower()],document)      
            p = para._p
            paragraph._p.addnext(p)
            
        elif (paragraph.text.strip()).lower in matching_heading: # check for a missed heading
            para = add_para(extraction_dictionary[paragraph.text.strip()])
            p = para._p            
            paragraph._p.addnext(p)
def Start_calling_from_here(filename,data_extracion,templete_heading):
    print("I am called")
    matching_heading,dic_heading = matching_heading_between_both(data_extraction_heading,templete_heading)
    doc = Document(filename)     
    move_para(doc,matching_heading,dic_heading) #sending data and document hanlde doc to function

    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    print(desktop)
    
    desktop = desktop + "\\"
    file_path = desktop + "Cv creation.docx"
    from pathlib import Path
    counter = 1
    while(True):
        my_file = Path(file_path)
        if my_file.is_file():
            file_path = desktop + "Cv Creation" + str(counter) +".docx" 
            print("File paath in if condition",file_path)
            counter = counter + 1
        else:
            doc.save(file_path)
            break
    print(file_path)
    return(file_path)
    
###############################################################################

rootname = Tk()
rootname.configure(background='black')
rootname.geometry("700x300")

def working_started():
    rootname.destroy()
    root1=Tk()
    root1.title("Cv style")
    root1.configure(background='black')
    root1.geometry("400x500")

    myFont_label = font.Font(size=15)
    myFont_button = font.Font(size=12)
    
    Heading_label =Label(root1,text="Curriculum Vitae Based Application",fg="white",bg="Black")
    Heading_label['font'] = myFont_label
    Heading_label.pack()

    Data_extraction_button = Button(root1,text="Select Cv for Data",fg="white",bg="Green",width=20,
                                    height=3,command=getting_path_for_extraction)
    Data_extraction_button['font'] = myFont_button
    Data_extraction_button.pack(pady=30)

    Cv_templete = Button(root1,text="Select Cv for Data",fg="white",bg="Green",width=20,
                                    height=3,command=getting_path_for_templete)
    Cv_templete['font'] = myFont_button
    Cv_templete.pack(pady=20)


    Heading_length_getting = Button(root1,text="Change Heading length",command=Clickme,fg="white",bg="Green",width=20,height=3)
    Heading_length_getting['font'] = myFont_button
    Heading_length_getting.pack(pady=20)

    Stated = Button(root1,text="Get Stated",fg="black",bg="Red",width=20,height=3,command=calling_function)
    Stated['font'] = myFont_button
    Stated.pack(pady=20)
    root.mainloop()

def Click_to_get1(horizontal,top):
    global heading_length_templete
    heading_length_templete =horizontal.get()
    messagebox.showinfo("Lenght of Heading is increase",'Length of heading is' + str(heading_length_templete))

def Click_to_get(horizontal,top):
    global heading_length_extraction
    heading_length_extraction =horizontal.get()
    messagebox.showinfo("Lenght of Heading is increase",'Length of heading is' + str(heading_length_templete))

def Clickme():
    top = Toplevel()
    top.geometry("400x300")
    top.title("Heading Length")
    top.configure(background='black')
    label = Label(top,text="Changing Heading Length",fg="white",bg="Black")
    label['font'] = myFont_label
    label.pack(pady=10)
    horizontal =Scale(top,from_=1,to=5,orient=HORIZONTAL)
    horizontal.pack()
    button = Button(top,text="Change Heading length of data Getting",command=lambda:Click_to_get(horizontal,top),fg="white",bg="Green")
    button.pack(pady=20)
    horizontal1 =Scale(top,from_=1,to=5,orient=HORIZONTAL)
    
    horizontal1.pack()
    button2 = Button(top,text="Change Heading length of templete",command=lambda:Click_to_get1(horizontal1,top),fg="white",bg="Green")
    button2.pack(pady=20)
    
def getting_path_for_templete():
    global templete
    global templete_heading
    global templete_dictionary
    Frame.FileTemplete= filedialog.askopenfilename(initialdir="/img",title="Select A file",filetypes=(("docx files","*.docx"),("all file","*.*")))
    templete= Frame.FileTemplete
    print("The value of temple is ",templete)

    templete_heading,templete_dictionary = main_function(templete,heading_length_templete)
    
    
def getting_path_for_extraction():
    print("Hello world")
    global Extraction
    global data_extraction_heading
    global extraction_dictionary
    Frame.FileExtraction= filedialog.askopenfilename(initialdir="/img",title="Select A file",filetypes=(("docx files","*.docx"),("all file","*.*")))
    Extraction=Frame.FileExtraction
    data_extraction_heading,extraction_dictionary = main_function(Extraction,heading_length_extraction)
    if Extraction != "" :
#         root.Cv_templete['state'] ='normal'
#         root.Cv_templete.pack()
        print("Hello")
def calling_function():
    global templete
    global data_extraction_heading
    global extraction_dictionary
    global templete_heading
    global templete_dictionary
    if Extraction == "" or templete == "":
        messagebox.showinfo("File Probelm","One file is missing either Data extraction File or Templete File")
    else:
        data_extraction_heading,extraction_dictionary = main_function(Extraction,heading_length_extraction)
        templete_heading,templete_dictionary = main_function(templete,heading_length_templete)
        file_path=Start_calling_from_here(templete,data_extraction_heading,templete_heading)
        messagebox.showinfo("File path where file is saved",file_path)
    



myFont_label = font.Font(size=15)
myFont_button = font.Font(size=12)
    
Heading_label =Label(rootname,text="Rules of Curriculum Vitae Based Application",fg="Red",bg="Black")
Heading_label['font'] = myFont_label
Heading_label.pack()

Rule1 = Label(rootname,text="1) If you want to change the heading lengths.Do it before selecting files",fg="white",bg="Black")
Rule1['font'] = myFont_label
Rule1.pack()

Rule2 = Label(rootname,text="2)If extra Heading came decreses the length of heading" ,fg="white",bg="Black")
Rule2['font'] = myFont_label
Rule2.pack()

Rule3 = Label(rootname,text="3)If some heading is missing increase the length of heading",fg="white",bg="Black")
Rule3['font'] = myFont_label
Rule3.pack()

Rule4 = Label(rootname,text="4) By heading extra and some means if it present in the data extraction file " ,fg="white",bg="Black")
Rule4['font'] = myFont_label
Rule4.pack()

Rule4 = Label(rootname,text=" and not show in result file then these functionality is used" ,fg="white",bg="Black")
Rule4['font'] = myFont_label
Rule4.pack()

Rule4 = Label(rootname,text="5)File is save on you Desktop with name Cv Creation" ,fg="white",bg="Black")
Rule4['font'] = myFont_label
Rule4.pack()


Stated1 = Button(rootname,text="ok,Get Stated",fg="black",bg="Green",width=20,height=2,command=working_started)
Stated1['font'] = myFont_button
Stated1.pack(pady=20)
rootname.mainloop()
